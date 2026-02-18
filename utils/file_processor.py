# utils/file_processor.py
import os
import logging
from typing import Dict, List, Optional
import PyPDF2
import pdfplumber
from docx import Document
import re

logger = logging.getLogger(__name__)

class FileProcessor:
    def __init__(self):
        self.supported_formats = ['.pdf', '.ppt', '.pptx', '.docx', '.txt']
    
    def extract_ppt_content(self, file_path: str) -> str:
        """Extract text content from PowerPoint files"""
        try:
            # Note: For PPT/PPTX files, we'll use a simplified approach
            # In production, you might want to use python-pptx library
            logger.warning(f"PPT extraction not fully implemented for {file_path}")
            
            # For now, return a placeholder message
            filename = os.path.basename(file_path)
            return f"Presentation file: {filename}. Please speak about the topics in your presentation."
            
        except Exception as e:
            logger.error(f"PPT extraction error: {str(e)}")
            return f"Error processing presentation file: {str(e)}"
    
    def extract_pdf_content(self, file_path: str) -> str:
        """Extract text content from PDF files"""
        try:
            content_parts = []
            
            # Try pdfplumber first (better for text extraction)
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page_num, page in enumerate(pdf.pages, 1):
                        text = page.extract_text()
                        if text and text.strip():
                            content_parts.append(f"Page {page_num}:\n{text.strip()}")
            except Exception as e:
                logger.warning(f"pdfplumber failed, trying PyPDF2: {str(e)}")
            
            # Fallback to PyPDF2
            if not content_parts:
                try:
                    with open(file_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        for page_num, page in enumerate(pdf_reader.pages, 1):
                            text = page.extract_text()
                            if text and text.strip():
                                content_parts.append(f"Page {page_num}:\n{text.strip()}")
                except Exception as e:
                    logger.error(f"PyPDF2 also failed: {str(e)}")
            
            if not content_parts:
                return "No extractable text found in PDF file."
            
            full_content = "\n\n".join(content_parts)
            logger.info(f"Extracted {len(full_content)} characters from PDF")
            return self._clean_extracted_content(full_content)
            
        except Exception as e:
            logger.error(f"PDF extraction error: {str(e)}")
            return f"Error processing PDF file: {str(e)}"
    
    def extract_docx_content(self, file_path: str) -> str:
        """Extract text content from Word documents"""
        try:
            doc = Document(file_path)
            content_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            content_parts.append(cell.text)
            
            full_content = "\n".join(content_parts)
            logger.info(f"Extracted {len(full_content)} characters from DOCX")
            return self._clean_extracted_content(full_content)
            
        except Exception as e:
            logger.error(f"DOCX extraction error: {str(e)}")
            return f"Error processing Word document: {str(e)}"
    
    def extract_txt_content(self, file_path: str) -> str:
        """Extract content from text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            logger.info(f"Extracted {len(content)} characters from TXT")
            return self._clean_extracted_content(content)
            
        except UnicodeDecodeError:
            # Try different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    content = file.read()
                return self._clean_extracted_content(content)
            except Exception as e:
                logger.error(f"TXT extraction error with alternative encoding: {str(e)}")
                return f"Error processing text file: {str(e)}"
        except Exception as e:
            logger.error(f"TXT extraction error: {str(e)}")
            return f"Error processing text file: {str(e)}"
    
    def extract_content(self, file_path: str) -> str:
        """Extract content from any supported file type"""
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            return self.extract_pdf_content(file_path)
        elif file_ext in ['.ppt', '.pptx']:
            return self.extract_ppt_content(file_path)
        elif file_ext == '.docx':
            return self.extract_docx_content(file_path)
        elif file_ext == '.txt':
            return self.extract_txt_content(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    def _clean_extracted_content(self, content: str) -> str:
        """Clean and normalize extracted content"""
        if not content:
            return ""
        
        # Remove extra whitespace
        content = re.sub(r'\s+', ' ', content)
        
        # Remove common file artifacts
        artifacts = [
            r'\x00',  # Null characters
            r'\ufeff',  # BOM
            r'\x0c',  # Form feed
        ]
        
        for artifact in artifacts:
            content = re.sub(artifact, '', content)
        
        # Limit content length to prevent overwhelming the system
        max_length = 10000
        if len(content) > max_length:
            content = content[:max_length] + "... [content truncated]"
        
        return content.strip()
    
    def get_file_info(self, file_path: str) -> Dict:
        """Get information about the uploaded file"""
        try:
            file_stats = os.stat(file_path)
            file_ext = os.path.splitext(file_path)[1].lower()
            file_size_kb = file_stats.st_size / 1024
            
            return {
                'filename': os.path.basename(file_path),
                'extension': file_ext,
                'size_kb': round(file_size_kb, 2),
                'supported': file_ext in self.supported_formats
            }
            
        except Exception as e:
            logger.error(f"File info extraction error: {str(e)}")
            return {
                'filename': os.path.basename(file_path),
                'extension': 'unknown',
                'size_kb': 0,
                'supported': False
            }
    
    def validate_file(self, file_path: str) -> Dict:
        """Validate uploaded file"""
        file_info = self.get_file_info(file_path)
        
        if not file_info['supported']:
            return {
                'valid': False,
                'error': f"Unsupported file format: {file_info['extension']}"
            }
        
        # Check file size (max 50MB)
        max_size_mb = 50
        if file_info['size_kb'] > max_size_mb * 1024:
            return {
                'valid': False,
                'error': f"File too large. Maximum size is {max_size_mb}MB"
            }
        
        return {
            'valid': True,
            'file_info': file_info
        }